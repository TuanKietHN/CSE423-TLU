from docx import Document

# Khởi tạo tài liệu
doc = Document()
doc.add_heading("Phân tích: Vì sao 9 luồng chậm hơn 2 luồng", level=1)

doc.add_paragraph("Bạn đang gặp tình huống 9 luồng chạy chậm hơn 2 luồng, dù đã \"tách thời gian khởi tạo luồng\" khỏi phần đo. "
                  "Có thể nguyên nhân không nằm ở đo thời gian nữa, mà do một số vấn đề khác ảnh hưởng đến hiệu năng thực tế. Cùng phân tích kỹ:")

doc.add_heading("✅ Đã Bỏ Tính Thời Gian Khởi Tạo?", level=2)
doc.add_paragraph("Rồi. Trong đoạn mã đã sửa:\n\n"
                  "- Các Thread được tạo xong trước khi gọi System.nanoTime().\n"
                  "- Đo thời gian bắt đầu sau khi khởi tạo luồng, nhưng trước khi start() → ✅ chính xác.")

doc.add_paragraph("Ví dụ trong mã:")
doc.add_paragraph("""
List<Thread> threads = new ArrayList<>();
// tạo thread trước

long startTime = System.nanoTime(); // ✅ Đúng chỗ
for (Thread t : threads) {
    t.start(); // chỉ chạy sau khi startTime đã ghi
}
""", style='Intense Quote')

doc.add_heading("❌ Vậy Tại Sao 9 Luồng Chạy Lâu Hơn?", level=2)

doc.add_heading("🔸 1. Số lượng luồng quá lớn so với số phần tử", level=3)
doc.add_paragraph("Bạn chia array.length thành k phần. Nếu mảng nhỏ (ví dụ 20 phần tử), chia thành 9 luồng nghĩa là mỗi luồng xử lý 2–3 phần tử "
                  "→ chi phí tạo luồng, đồng bộ (synchronized, wait, notifyAll) tốn nhiều hơn thời gian xử lý thực tế.")

doc.add_heading("🔸 2. synchronized + wait(10) làm nghẽn", level=3)
doc.add_paragraph("Mỗi khi tìm được số chính phương, các luồng:\n"
                  "- synchronized(lock)\n"
                  "- notifyAll() rồi wait(10) → tất cả luồng phải tranh nhau vào vùng synchronized, khiến các luồng chờ nhau liên tục.\n"
                  "→ Với nhiều luồng, nghẽn lock càng nặng → tốc độ kém đi.")

doc.add_paragraph("👉 Càng nhiều luồng, vùng đồng bộ càng trở thành “nút thắt cổ chai” (bottleneck).")

doc.add_heading("✅ Cách khắc phục:", level=2)

doc.add_heading("✔️ Cách 1: Bỏ synchronized, wait, notifyAll", level=3)
doc.add_paragraph("Chỉ cần in ra kết quả là đủ, không cần đồng bộ khi không thay đổi tài nguyên dùng chung (trừ result, đã là CopyOnWriteArrayList an toàn rồi).")

doc.add_heading("✔️ Cách 2: Tăng kích thước mảng", level=3)
doc.add_paragraph("Mảng 20–30 phần tử không đủ để đánh giá tốc độ đa luồng. Bạn nên thử với mảng 100.000 hoặc 1.000.000 phần tử.")

doc.add_heading("🔧 Gợi ý sửa findPerfectSquaresWithKThreads() để benchmark chuẩn hơn:", level=2)
doc.add_paragraph("""
private static List<PerfectSquare> findPerfectSquaresWithKThreads(int[] array, int k) {
    var result = new CopyOnWriteArrayList<PerfectSquare>();
    var latch = new CountDownLatch(k);
    int chunkSize = array.length / k;

    List<Thread> threads = new ArrayList<>();

    for (int t = 0; t < k; t++) {
        final int start = t * chunkSize;
        final int end = (t == k - 1) ? array.length : (t + 1) * chunkSize;

        Thread thread = new Thread(() -> {
            for (int i = start; i < end; i++) {
                if (isPerfectSquare(array[i])) {
                    result.add(new PerfectSquare(i, array[i]));
                    // Bỏ synchronized để tránh nghẽn luồng
                    // System.out.println("Luong " + Thread.currentThread().getId() + ": " + array[i]);
                }
            }
            latch.countDown();
        });
        threads.add(thread);
    }

    long startTime = System.nanoTime();

    for (Thread t : threads) {
        t.start();
    }

    try {
        latch.await();
    } catch (InterruptedException e) {
        Thread.currentThread().interrupt();
    }

    long endTime = System.nanoTime();
    System.out.println("Thoi gian thuc thi: " + (endTime - startTime) / 1_000_000.0 + " ms");

    return result;
}
""", style='Intense Quote')

doc.add_heading("✅ Kết luận:", level=2)

doc.add_paragraph("""
| Nguyên nhân khiến 9 luồng chậm hơn | Cách khắc phục |
|------------------------------------|----------------|
| Mảng quá nhỏ → overhead chia nhỏ | Dùng mảng lớn hơn |
| Gây nghẽn do synchronized + wait() | Loại bỏ vùng đồng bộ khi không cần |
| Thời gian đo đã chính xác | Không cần sửa nữa |
""", style='Intense Quote')

# Lưu file
doc.save("phan_tich_da_luong.docx")
